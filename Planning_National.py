import streamlit as st
from supabase import create_client, Client
import pandas as pd
import calendar
import json
import os
import io
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Identifiants Supabase (à remplacer par les vôtres)
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]


# Initialisation du client
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# --- INITIALISATION DU SESSION STATE ---
if 'config' not in st.session_state:
    donnees_initiales = charger_donnees()
    if donnees_initiales:
        st.session_state.config = donnees_initiales['config']
        st.session_state.composition = donnees_initiales['composition']
        st.session_state.conges = donnees_initiales['conges']
        st.session_state.liste_unites = donnees_initiales['liste_unites']
        st.session_state.base_agents = donnees_initiales['base_agents']
    else:
        # Valeurs de secours si Supabase est vide ou inaccessible
        st.session_state.config = {"region": "Centre-Ouest", "unite_active": "Unité de Soins"}
        st.session_state.composition = {}
        st.session_state.conges = []
        st.session_state.liste_unites = ["Unité de Soins"]
        st.session_state.base_agents = []

MOIS_FR = {
    1: "JANVIER", 2: "FÉVRIER", 3: "MARS", 4: "AVRIL", 
    5: "MAI", 6: "JUIN", 7: "JUILLET", 8: "AOÛT", 
    9: "SEPTEMBRE", 10: "OCTOBRE", 11: "NOVEMBRE", 12: "DÉCEMBRE"
}

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="Validateur National Santé", layout="wide")
#FILE_PATH = "config_national_final.json"

# --- 2. FONCTIONS DE GESTION DES DONNÉES ---

def charger_donnees():
    try:
        # 1. Récupération de la configuration (ligne avec ID 1)
        res_conf = supabase.table("configuration").select("*").eq("id", 1).execute()
        
        # 2. Récupération de la liste complète des agents
        res_agents = supabase.table("base_agents").select("*").execute()
        
        data = {
            'config': {}, 
            'composition': {}, 
            'conges': [], 
            'liste_unites': ["Unité de Soins"],
            'base_agents': []
        }

        # Traitement de la configuration
        if res_conf.data:
            conf = res_conf.data[0]
            data['config'] = conf.get('config_globale', {})
            data['composition'] = conf.get('composition_equipes', {})
            data['liste_unites'] = conf.get('liste_unites', ["Unité de Soins"])
            
            # Conversion des dates de congés (format texte -> format date Python)
            conges_bruts = conf.get('conges', [])
            for c in conges_bruts:
                c['debut'] = datetime.strptime(c['debut'], "%Y-%m-%d").date()
                c['fin'] = datetime.strptime(c['fin'], "%Y-%m-%d").date()
            data['conges'] = conges_bruts
            
        # Traitement des agents
        if res_agents.data:
            data['base_agents'] = res_agents.data
            
        return data
    except Exception as e:
        st.error(f"Erreur de chargement Cloud : {e}")
        return None

def sauvegarder_donnees():
    try:
        # 1. Sauvegarde des agents (Gestion du conflit sur le NOM pour éviter les doublons)
        for agent in st.session_state.get('base_agents', []):
            agent_payload = {
                "nom": agent["nom"],
                "emploi": agent.get("emploi", ""),
                "matricule": agent.get("matricule", "")
            }
            supabase.table("base_agents").upsert(agent_payload, on_conflict="nom").execute()

        # 2. Sauvegarde de la configuration globale
        config_payload = {
            "id": 1,
            "liste_unites": st.session_state.get('liste_unites', ["Unité de Soins"]),
            "config_globale": st.session_state.get('config', {}),
            "composition_equipes": st.session_state.get('composition', {}),
            "conges": [{**c, "debut": str(c['debut']), "fin": str(c['fin'])} 
                       for c in st.session_state.get('conges', [])]
        }
        supabase.table("configuration").upsert(config_payload).execute()
        
    except Exception as e:
        st.error(f"Erreur de sauvegarde Cloud : {e}")

# =========================================================
# 3...ALGORITHME DE PLANNING NATIONAL SANTÉ - VERSION FINALE (CORRIGÉE V13)
# =========================================================
# --LE COEUR : CALCUL ET LISSAGE RÉEL ---
def calculer_planning_pro(annee, mois, nb_equipes, reliquat_semaine_derniere=10):
    jours_dans_mois = calendar.monthrange(annee, mois)[1]
    u_active = st.session_state.config.get('unite_active', 'Maternité')
    equipes_dict = {e_id: st.session_state.composition.get(f"{u_active}_{e_id}", []) for e_id in range(1, nb_equipes + 1)}
    set_conges = {(c['agent'], d) for c in st.session_state.conges for d in [c['debut'] + timedelta(days=x) for x in range((c['fin'] - c['debut']).days + 1)]}

    planning_brut = {}
    heures_hebdo = {}
    curseur_rotation = 0 
    historique_ids_g = {} 

    # --- PASSE 1 : ATTRIBUTION ET COMPTAGE BRUT ---
    for j in range(1, jours_dans_mois + 1):
        dt = date(annee, mois, j)
        sem_key = f"Semaine {dt.isocalendar()[1]}"
        is_we = dt.weekday() >= 5 

        # Détermination de la garde
        id_g_reel = None
        tentatives = 0
        while tentatives < nb_equipes:
            id_test = (curseur_rotation % nb_equipes) + 1
            membres = equipes_dict.get(id_test, [])
            if membres and any((m, dt) not in set_conges for m in membres):
                id_g_reel = id_test
                break
            curseur_rotation += 1
            tentatives += 1
        
        if id_g_reel is None: id_g_reel = (curseur_rotation % nb_equipes) + 1
        historique_ids_g[j] = id_g_reel
        curseur_rotation += 1 

        id_r_reel = historique_ids_g.get(j-1)
        id_j2_reel = historique_ids_g.get(j-2)

        for e_id, membres in equipes_dict.items():
            for n in membres:
                if n not in heures_hebdo: heures_hebdo[n] = {}
                if sem_key not in heures_hebdo[n]:
                    heures_hebdo[n][sem_key] = reliquat_semaine_derniere if j <= 7 else 0

                res = {"type": "Repos", "heures": 0}
                
                if (n, dt) in set_conges:
                    res = {"type": "Congé", "heures": 0}
                elif e_id == id_g_reel:
                    res = {"type": "Garde", "heures": 0}
                elif e_id == id_r_reel:
                    res = {"type": "Repos", "heures": 0}
                elif e_id == id_j2_reel:
                    # REPRISE OBLIGATOIRE J+2 (10h protégée par le flag j2)
                    res = {"type": "Reprise J+2", "heures": 10, "j2": True} 
                    heures_hebdo[n][sem_key] += 10
                elif not is_we:
                    res = {"type": "Journée", "heures": 10}
                    heures_hebdo[n][sem_key] += 10
                elif is_we:
                    # On met 10h d'office ; le lissage Passe 2 décidera de réduire à 5h
                    res = {"type": "Week-end", "heures": 10}
                    heures_hebdo[n][sem_key] += 10

                planning_brut[(n, j)] = res

    # --- PASSE 2 : LISSAGE RÉTROACTIF (MÉTACOGNITION) ---
    for n in heures_hebdo.keys():
        for sem_key, total in heures_hebdo[n].items():
            if total > 40:
                # On scanne les jours de la semaine à l'envers (Dimanche vers Lundi)
                for j_corr in range(jours_dans_mois, 0, -1):
                    dt_c = date(annee, mois, j_corr)
                    if f"Semaine {dt_c.isocalendar()[1]}" == sem_key:
                        target = planning_brut.get((n, j_corr))
                        
                        # Sécurité : On ne touche qu'aux jours de 10h qui ne sont pas des reprises J+2
                        if target and target.get("heures") == 10 and not target.get("j2"):
                            if target["type"] in ["Journée", "Week-end"]:
                                if target["type"] == "Journée":
                                    target["type"] = "Demi-journée"
                            
                                target["heures"] = 5
                                heures_hebdo[n][sem_key] -= 5
                            
                                if heures_hebdo[n][sem_key] <= 40:
                                    break
    return planning_brut, heures_hebdo

# --- 5. EXPORT WORD (INTACT) ---
def exporter_vers_word(recap_data, config, mois_num, annee_num):
    doc = Document()
    section = doc.sections[-1]; section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    header_table = doc.add_table(rows=1, cols=2)
    c1, c2 = header_table.rows[0].cells
    c1.text = f"MINISTERE DE LA SANTE\n-----------------\nREGION DU {config.get('region','')}\n-----------------\nDISTRICT DE {config.get('district','')}\n-----------------\n{config.get('nom_csps','')}"
    c2.text = "BURKINA FASO\n-----------------\nLa Patrie ou la mort, Nous vaincrons"
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    titre = doc.add_paragraph(f"\nPROGRAMME DE TRAVAIL DU MOIS DE {MOIS_FR[mois_num]} {annee_num} ({config.get('unite_active','').upper()})")
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER; titre.runs[0].bold = True
    table = doc.add_table(rows=2, cols=9); table.style = 'Table Grid'
    h1 = table.rows[0].cells; h1[0].text = "Eq"; h1[1].text = "Nom & Prénoms"; h1[2].text = "Emploi"; h1[3].text = "Matricule"
    h1[4].merge(h1[5]).merge(h1[6]); h1[4].text = "Travail de jour"; h1[7].text = "Garde"; h1[8].text = "Repos"
    h2 = table.rows[1].cells; h2[4].text = "7h30-17h30"; h2[5].text = "7h30-12h30"; h2[6].text = "Week-end"
    for i in [0,1,2,3,7,8]: table.cell(0,i).merge(table.cell(1,i))
    for entry in recap_data:
        membres = entry['Membres'].split(" / ")
        first_row_idx = len(table.rows)
        for idx, nom_complet in enumerate(membres):
            row = table.add_row().cells
            nom_nettoye = nom_complet.replace(" (🏖️)", "").strip()
            agent_info = next((a for a in st.session_state.base_agents if a['nom'].strip() == nom_nettoye), {})
            row[1].text = nom_nettoye; row[2].text = agent_info.get('emploi', ''); row[3].text = agent_info.get('matricule', '')
            if idx == 0:
                row[0].text = str(entry['N°']); row[4].text = entry.get('Journée', ''); row[5].text = entry.get('Demi-journée', '')
                row[6].text = entry.get('Week-end et fériés', ''); row[7].text = entry.get('Garde', ''); row[8].text = entry.get('Repos', '')
        if len(membres) > 1:
            for col_idx in [0, 4, 5, 6, 7, 8]: table.cell(first_row_idx, col_idx).merge(table.cell(len(table.rows)-1, col_idx))
    doc.add_paragraph("\n"); sig_table = doc.add_table(rows=1, cols=2); s1, s2 = sig_table.rows[0].cells
    s1.text = f"L’Infirmier chef de poste\n\n\n{config.get('nom_icp', 'ICP')}"
    s2.text = f"Le Médecin Chef du District\n\n\n{config.get('nom_mcd', 'MCD')}"
    s2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    target = io.BytesIO(); doc.save(target); target.seek(0)
    return target

# --- 6. INTERFACE ET AFFICHAGE ---
with st.sidebar:
    st.header("⚙️ Configuration")
    with st.expander("🏥 Gérer les Unités"):
        nouvelle = st.text_input("Nouvelle unité")
        if st.button("➕ Ajouter") and nouvelle:
            if nouvelle not in st.session_state.liste_unites:
                st.session_state.liste_unites.append(nouvelle); sauvegarder_donnees(); st.rerun()
    st.session_state.config["region"] = st.text_input("Région", value=st.session_state.config.get("region", "Centre-Ouest"))
    st.session_state.config["district"] = st.text_input("District", value=st.session_state.config.get("district", "KOUDOUGOU"))
    st.session_state.config["nom_csps"] = st.text_input("CSPS", value=st.session_state.config.get("nom_csps", "SECTEUR 10"))
    st.session_state.config["nom_icp"] = st.text_input("Nom ICP", value=st.session_state.config.get("nom_icp", ""))
    st.session_state.config["nom_mcd"] = st.text_input("Nom MCD", value=st.session_state.config.get("nom_mcd", ""))
    st.divider()
    u_active = st.selectbox("Unité Active", st.session_state.liste_unites, index=st.session_state.liste_unites.index(st.session_state.config.get("unite_active", st.session_state.liste_unites[0])))
    st.session_state.config["unite_active"] = u_active
    nb_eq = st.slider("Nombre d'Équipes", 2, 15, value=st.session_state.config.get("nb_equipes", 10))
    st.session_state.config["nb_equipes"] = nb_eq
    val_reliquat = st.number_input("Reliquat (Heures mois précédent)", min_value=0, max_value=40, value=st.session_state.config.get("reliquat", 10))
    st.session_state.config["reliquat"] = val_reliquat
    annee = st.number_input("Année", value=2026)
    mois = st.slider("Mois", 1, 12, value=4)
    if st.button("💾 Sauvegarder Config"): sauvegarder_donnees(); st.success("Enregistré !")

# --- EXÉCUTION DU CALCUL ---
planning_final, heures_hebdo = calculer_planning_pro(annee, mois, nb_eq, val_reliquat)

t1, t2, t3, t4, t5 = st.tabs(["✅ Validation", "📊 Vue Regroupée", "👥 Équipes", "🏖️ Congés", "🗂️ Base Agents"])

with t1:
    for nom, semaines in heures_hebdo.items():
        with st.expander(f"👤 {nom}"):
            cols = st.columns(len(semaines))
            for i, (s, total) in enumerate(semaines.items()):
                if 35 <= total <= 40: cols[i].success(f"{s}: {total}h")
                elif total > 40: cols[i].error(f"{s}: {total}h")
                else: cols[i].warning(f"{s}: {total}h")

with t2:
    recap = []
    for i in range(1, nb_eq + 1):
        key_compo = f"{u_active}_{i}"
        membres_bruts = st.session_state.composition.get(key_compo, [])
        jrs = {"Journée": [], "Demi-journée": [], "Week-end": [], "Garde": [], "Repos": []}
        if membres_bruts:
            for d in range(1, calendar.monthrange(annee, mois)[1] + 1):
                shift = {}
                for agent in membres_bruts:
                    s = planning_final.get((agent, d), {})
                    if s.get("type") != "Congé": shift = s; break 
                t = shift.get("type"); h = shift.get("heures", 0)
                if t == "Week-end": jrs["Week-end"].append(str(d))
                elif t == "Garde": jrs["Garde"].append(str(d))
                elif t == "Repos": jrs["Repos"].append(str(d))
                elif t == "Reprise J+2": jrs["Journée"].append(str(d))
                elif t == "Journée": jrs["Journée"].append(str(d))
                elif t == "Demi-journée": jrs["Demi-journée"].append(str(d))
        noms_equipe = [f"{m} (🏖️)" if any(c['agent']==m for c in st.session_state.conges) else m for m in membres_bruts]
        recap.append({"N°": i, "Membres": " / ".join(noms_equipe), "Journée": ", ".join(jrs["Journée"]), "Demi-journée": ", ".join(jrs["Demi-journée"]), "Week-end et fériés": ", ".join(jrs["Week-end"]), "Garde": ", ".join(jrs["Garde"]), "Repos": ", ".join(jrs["Repos"])})
    if recap:
        st.table(pd.DataFrame(recap))
        st.download_button(label="📄 Exporter vers WORD", data=exporter_vers_word(recap, st.session_state.config, mois, annee), file_name=f"Planning_{u_active}_{MOIS_FR[mois]}.docx")

# Onglets d'administration (Équipes, Congés, Base Agents) inchangés pour préserver ton flux de travail
with t3:
    st.subheader("👥 Équipes")
    options_agents = [a["nom"] for a in st.session_state.base_agents]
    tous_affectes = [x for eid in range(1, nb_eq + 1) for x in st.session_state.composition.get(f"{u_active}_{eid}", [])]
    cols = st.columns(3)
    for i in range(1, nb_eq + 1):
        key = f"{u_active}_{i}"; anciens = st.session_state.composition.get(key, [])
        dispo = [a for a in options_agents if a not in [x for x in tous_affectes if x not in anciens]]
        res = cols[(i-1)%3].multiselect(f"Équipe {i}", options=dispo, default=anciens, key=f"sel_{key}")
        if res != anciens: st.session_state.composition[key] = res; sauvegarder_donnees(); st.rerun()

with t4:
    with st.form("f_conge"):
        c_nom = st.selectbox("Agent", [a["nom"] for a in st.session_state.base_agents])
        d1, d2 = st.columns(2); b = d1.date_input("Début"); f = d2.date_input("Fin")
        if st.form_submit_button("Enregistrer"): st.session_state.conges.append({"agent": c_nom, "debut": b, "fin": f}); sauvegarder_donnees(); st.rerun()
    if st.session_state.conges: st.table(pd.DataFrame(st.session_state.conges))

with t5:
    with st.form("new_agent"):
        n = st.text_input("Nom & Prénoms"); e = st.text_input("Emploi"); m = st.text_input("Matricule")
        if st.form_submit_button("Ajouter") and n:
            if not any(a['nom']==n for a in st.session_state.base_agents):
                st.session_state.base_agents.append({"nom":n,"emploi":e,"matricule":m}); sauvegarder_donnees(); st.rerun()
    st.dataframe(pd.DataFrame(st.session_state.base_agents))



